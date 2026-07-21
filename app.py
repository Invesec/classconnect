# CRITICAL: this must be the very first thing that runs in this file, before
# any other import. If Render's deployment is using an eventlet worker
# (gunicorn --worker-class eventlet), eventlet needs to monkey-patch Python's
# threading/socket/ssl internals before ANY other module (Flask, SQLAlchemy,
# sqlite3, etc.) creates lock objects — otherwise those locks stay as
# regular OS locks that eventlet's greenlets can't coordinate with, which is
# exactly the "RuntimeError: cannot notify on un-acquired lock" /
# "N RLock(s) were not greened" errors. Relying on gunicorn's own worker
# startup to do this at the right time is fragile depending on preload
# settings, so we do it ourselves, unconditionally, as early as possible.
# If eventlet isn't installed (e.g. running via plain `python app.py` in
# development), this is a harmless no-op.
try:
    import eventlet
    eventlet.monkey_patch()
except ImportError:
    pass

import os, secrets, string, random
from datetime import datetime, timezone, timedelta

import requests
from flask import (Flask, render_template, redirect, url_for, flash,
                   request, abort, send_from_directory, jsonify, session)
from flask_sqlalchemy import SQLAlchemy
from flask_login import (LoginManager, UserMixin, login_user, logout_user,
                         login_required, current_user)
from flask_mail import Mail, Message
from flask_socketio import SocketIO, join_room, leave_room, emit
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document as DocxDocument
from docx.shared import Pt, Inches

basedir = os.path.abspath(os.path.dirname(__file__))

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')

# ── Database ──────────────────────────────────────────────────────────────────
# DATABASE_URL, if set, points at a real Postgres database (e.g. Neon's free
# tier) — this is what actually persists data across Render restarts/redeploys.
# Falls back to a local SQLite file if unset (fine for local development, but
# NOT persistent on Render's free tier — every restart wipes it).
_db_url = os.environ.get(
    'DATABASE_URL',
    f"sqlite:///{os.path.join(basedir, 'instance', 'classconnect.db')}"
)
# Some providers (Neon, Heroku, etc.) hand out connection strings starting
# with "postgres://", which older/newer SQLAlchemy versions reject — it
# requires the "postgresql://" scheme instead. Normalize it here so pasting
# a connection string straight from the provider's dashboard just works.
if _db_url.startswith('postgres://'):
    _db_url = _db_url.replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_DATABASE_URI'] = _db_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# ── File uploads ──────────────────────────────────────────────────────────────
app.config['UPLOAD_FOLDER'] = os.path.join(basedir, 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024   # 25 MB

# ── Email (Flask-Mail) ────────────────────────────────────────────────────────
# Configure via .env or environment variables on your server.
# For Gmail: allow "App Passwords" in your Google account and set:
#   MAIL_USERNAME=you@gmail.com   MAIL_PASSWORD=your-app-password
app.config['MAIL_SERVER']   = os.environ.get('MAIL_SERVER',   'smtp.gmail.com')
app.config['MAIL_PORT']     = int(os.environ.get('MAIL_PORT', 587))
app.config['MAIL_USE_TLS']  = os.environ.get('MAIL_USE_TLS',  'true').lower() == 'true'
app.config['MAIL_USERNAME'] = os.environ.get('MAIL_USERNAME', '')
app.config['MAIL_PASSWORD'] = os.environ.get('MAIL_PASSWORD', '')
# IMPORTANT (deliverability): the "From" address MUST be the same mailbox that
# authenticates via MAIL_USERNAME/MAIL_PASSWORD. Sending through Gmail's SMTP
# while claiming a "From" on a different, unverified domain (e.g.
# noreply@classconnect.fuo) fails SPF/DKIM alignment and is a top reason mail
# lands straight in spam. Only the display NAME is customizable — the address
# always matches the authenticated account.
MAIL_SENDER_NAME = os.environ.get('MAIL_SENDER_NAME', 'ClassConnect')
app.config['MAIL_DEFAULT_SENDER'] = (MAIL_SENDER_NAME, app.config['MAIL_USERNAME']) \
    if app.config['MAIL_USERNAME'] else 'noreply@classconnect.fuo'

# Render's FREE tier blocks all outbound SMTP traffic (ports 25, 465, 587)
# as a platform-level restriction — this has nothing to do with credentials
# or code, and shows up as [Errno 110] ETIMEDOUT. If BREVO_API_KEY is set,
# email is sent over Brevo's HTTPS API instead (port 443, never blocked),
# which works on the free tier. If it's not set, falls back to normal SMTP
# via Flask-Mail (works fine on a paid Render plan, or elsewhere).
# Sign up free at https://www.brevo.com (300 emails/day free), verify a
# sender address there, then set BREVO_API_KEY + BREVO_SENDER_EMAIL.
app.config['BREVO_API_KEY']      = os.environ.get('BREVO_API_KEY', '')
app.config['BREVO_SENDER_EMAIL'] = os.environ.get('BREVO_SENDER_EMAIL', app.config['MAIL_USERNAME'])

ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'ppt', 'pptx', 'webm', 'mp4', 'mp3', 'wav', 'ogg'}
OTP_EXPIRY_MINUTES = 10
RESET_TOKEN_EXPIRY_MINUTES = 30

db       = SQLAlchemy(app)
mail     = Mail(app)
socketio = SocketIO(app, cors_allowed_origins='*', async_mode='threading')
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message_category = 'info'


# ── Models ────────────────────────────────────────────────────────────────────

class Organization(db.Model):
    """The top-level tenant. Every user belongs to exactly one. Started as
    university-only; this generalizes it to any organization (company,
    NGO, training provider, community group, etc.) while keeping the
    underlying role/course/session model identical — only the vocabulary
    shown to users changes based on org_type. See ORG_LABELS below."""
    __tablename__ = 'organizations'
    id          = db.Column(db.Integer, primary_key=True)
    name        = db.Column(db.String(150), nullable=False)
    org_type    = db.Column(db.String(20), nullable=False, default='university')  # 'university' | 'organization'
    join_code   = db.Column(db.String(10), unique=True, nullable=False)
    created_at  = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    members     = db.relationship('User', backref='organization', lazy=True)


# Vocabulary shown to users, based on their organization's type. The
# underlying role values in the database ('lecturer', 'student', 'admin')
# never change — only these display labels do — so none of the existing
# permission logic (role_required, is_course_lecturer, etc.) needed to
# change to support organizations.
ORG_LABELS = {
    'university': {
        'org_word': 'University', 'admin_role': 'Lecturer', 'admin_role_plural': 'Lecturers',
        'member_role': 'Student', 'member_role_plural': 'Students',
        'unit': 'Course', 'unit_plural': 'Courses', 'id_field': 'Matric Number',
    },
    'organization': {
        'org_word': 'Organization', 'admin_role': 'Facilitator', 'admin_role_plural': 'Facilitators',
        'member_role': 'Member', 'member_role_plural': 'Members',
        'unit': 'Program', 'unit_plural': 'Programs', 'id_field': 'Member ID',
    },
}


def get_org_labels(organization):
    org_type = organization.org_type if organization else 'university'
    return ORG_LABELS.get(org_type, ORG_LABELS['university'])


class User(UserMixin, db.Model):
    __tablename__ = 'users'
    id              = db.Column(db.Integer, primary_key=True)
    full_name       = db.Column(db.String(120), nullable=False)
    email           = db.Column(db.String(120), unique=True, nullable=False)
    password_hash   = db.Column(db.String(255), nullable=False)
    role            = db.Column(db.String(20), nullable=False)   # lecturer | student | admin
    id_number       = db.Column(db.String(30), unique=True, nullable=True)
    # Unique identity number: Matric Number for students, Staff/Lecturer ID for lecturers.
    organization_id = db.Column(db.Integer, db.ForeignKey('organizations.id'), nullable=True)
    email_verified  = db.Column(db.Boolean, default=False)
    otp             = db.Column(db.String(6))
    otp_expires     = db.Column(db.DateTime)
    reset_token     = db.Column(db.String(64))
    reset_token_expires = db.Column(db.DateTime)
    created_at      = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

    courses_taught = db.relationship('Course', backref='lecturer', lazy=True,
                                     foreign_keys='Course.lecturer_id')

    @property
    def labels(self):
        return get_org_labels(self.organization)

    def set_password(self, pw):
        self.password_hash = generate_password_hash(pw)

    def check_password(self, pw):
        return check_password_hash(self.password_hash, pw)

    def generate_otp(self):
        self.otp = ''.join([str(random.randint(0, 9)) for _ in range(6)])
        self.otp_expires = datetime.now(timezone.utc) + timedelta(minutes=OTP_EXPIRY_MINUTES)
        return self.otp

    def verify_otp(self, code):
        if not self.otp or not self.otp_expires:
            return False
        expires = self.otp_expires
        if expires.tzinfo is None:
            expires = expires.replace(tzinfo=timezone.utc)
        return self.otp == code and datetime.now(timezone.utc) < expires

    def generate_reset_token(self):
        self.reset_token = secrets.token_urlsafe(32)
        self.reset_token_expires = datetime.now(timezone.utc) + timedelta(minutes=RESET_TOKEN_EXPIRY_MINUTES)
        return self.reset_token

    def verify_reset_token(self, token):
        if not self.reset_token or not self.reset_token_expires or not token:
            return False
        expires = self.reset_token_expires
        if expires.tzinfo is None:
            expires = expires.replace(tzinfo=timezone.utc)
        return secrets.compare_digest(self.reset_token, token) and datetime.now(timezone.utc) < expires

    def clear_reset_token(self):
        self.reset_token = None
        self.reset_token_expires = None


class Course(db.Model):
    __tablename__ = 'courses'
    id                    = db.Column(db.Integer, primary_key=True)
    course_code           = db.Column(db.String(20), nullable=False)
    title                 = db.Column(db.String(150), nullable=False)
    description           = db.Column(db.Text)
    enrolment_code        = db.Column(db.String(10), unique=True, nullable=False)
    lecturer_invite_code  = db.Column(db.String(10), unique=True)
    lecturer_id           = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    created_at            = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

    enrolments   = db.relationship('Enrolment',        backref='course', lazy=True, cascade='all, delete-orphan')
    sessions     = db.relationship('ClassSession',      backref='course', lazy=True, cascade='all, delete-orphan')
    materials    = db.relationship('CourseMaterial',    backref='course', lazy=True, cascade='all, delete-orphan')
    links        = db.relationship('CourseLink',        backref='course', lazy=True, cascade='all, delete-orphan')
    co_lecturers = db.relationship('CourseCoLecturer',  backref='course', lazy=True, cascade='all, delete-orphan')
    assignments  = db.relationship('Assignment',        backref='course', lazy=True, cascade='all, delete-orphan')


class CourseCoLecturer(db.Model):
    """A lecturer other than the course owner who has been granted full
    lecturer access to this course (view/manage materials, run sessions,
    moderate live classes, create assignments) — everything except
    deleting the course itself, which stays with the original owner."""
    __tablename__ = 'course_co_lecturers'
    id           = db.Column(db.Integer, primary_key=True)
    course_id    = db.Column(db.Integer, db.ForeignKey('courses.id'), nullable=False)
    lecturer_id  = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    invited_at   = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    lecturer     = db.relationship('User', foreign_keys=[lecturer_id])

    __table_args__ = (db.UniqueConstraint('course_id', 'lecturer_id', name='uq_course_colecturer'),)


class Assignment(db.Model):
    __tablename__ = 'assignments'
    id           = db.Column(db.Integer, primary_key=True)
    course_id    = db.Column(db.Integer, db.ForeignKey('courses.id'), nullable=False)
    created_by   = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    title        = db.Column(db.String(200), nullable=False)
    instructions = db.Column(db.Text)
    due_date     = db.Column(db.DateTime)
    created_at   = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

    submissions = db.relationship('AssignmentSubmission', backref='assignment', lazy=True, cascade='all, delete-orphan')


class AssignmentSubmission(db.Model):
    __tablename__ = 'assignment_submissions'
    id               = db.Column(db.Integer, primary_key=True)
    assignment_id    = db.Column(db.Integer, db.ForeignKey('assignments.id'), nullable=False)
    student_id       = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    filename         = db.Column(db.String(255), nullable=False)
    stored_filename  = db.Column(db.String(255), nullable=False)
    submitted_at     = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    student          = db.relationship('User', foreign_keys=[student_id])

    __table_args__ = (db.UniqueConstraint('assignment_id', 'student_id', name='uq_one_submission_per_student'),)


class Enrolment(db.Model):
    __tablename__ = 'enrolments'
    id         = db.Column(db.Integer, primary_key=True)
    student_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    course_id  = db.Column(db.Integer, db.ForeignKey('courses.id'), nullable=False)
    enrolled_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    student    = db.relationship('User', foreign_keys=[student_id])
    __table_args__ = (db.UniqueConstraint('student_id', 'course_id', name='uq_student_course'),)


class ClassSession(db.Model):
    __tablename__ = 'class_sessions'
    id         = db.Column(db.Integer, primary_key=True)
    course_id  = db.Column(db.Integer, db.ForeignKey('courses.id'), nullable=False)
    topic      = db.Column(db.String(200))
    status     = db.Column(db.String(20), default='active')   # active | closed
    started_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    closed_at  = db.Column(db.DateTime)

    participants       = db.relationship('SessionParticipant', backref='session', lazy=True, cascade='all, delete-orphan')
    attendance_records = db.relationship('AttendanceRecord',   backref='session', lazy=True, cascade='all, delete-orphan')
    chat_messages       = db.relationship('SessionChatMessage', backref='session', lazy=True, cascade='all, delete-orphan')


class SessionChatMessage(db.Model):
    """Persisted live-session chat. Previously this was broadcast-only and
    vanished on refresh — now it's stored so re-entering an active session
    (or just reloading the page) shows the full history, and a message's
    author can edit it afterward if they made a typo."""
    __tablename__ = 'session_chat_messages'
    id          = db.Column(db.Integer, primary_key=True)
    session_id  = db.Column(db.Integer, db.ForeignKey('class_sessions.id'), nullable=False)
    user_id     = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    text        = db.Column(db.Text, nullable=False)
    created_at  = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    edited_at   = db.Column(db.DateTime)
    user        = db.relationship('User', foreign_keys=[user_id])


class SessionParticipant(db.Model):
    __tablename__ = 'session_participants'
    id              = db.Column(db.Integer, primary_key=True)
    session_id      = db.Column(db.Integer, db.ForeignKey('class_sessions.id'), nullable=False)
    student_id      = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    joined_at       = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    still_connected = db.Column(db.Boolean, default=True)
    last_seen       = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    student         = db.relationship('User', foreign_keys=[student_id])
    __table_args__  = (db.UniqueConstraint('session_id', 'student_id', name='uq_session_student'),)


class AttendanceRecord(db.Model):
    __tablename__ = 'attendance_records'
    id          = db.Column(db.Integer, primary_key=True)
    session_id  = db.Column(db.Integer, db.ForeignKey('class_sessions.id'), nullable=False)
    student_id  = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    status      = db.Column(db.String(20), default='present')
    recorded_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    student     = db.relationship('User', foreign_keys=[student_id])


class CourseMaterial(db.Model):
    __tablename__ = 'course_materials'
    id               = db.Column(db.Integer, primary_key=True)
    course_id        = db.Column(db.Integer, db.ForeignKey('courses.id'), nullable=False)
    uploader_id      = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    title            = db.Column(db.String(200), nullable=False)
    filename         = db.Column(db.String(255), nullable=False)
    stored_filename  = db.Column(db.String(255), nullable=False)
    upload_date      = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    uploader         = db.relationship('User', foreign_keys=[uploader_id])


class CourseLink(db.Model):
    __tablename__ = 'course_links'
    id           = db.Column(db.Integer, primary_key=True)
    course_id    = db.Column(db.Integer, db.ForeignKey('courses.id'), nullable=False)
    added_by_id  = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    title        = db.Column(db.String(200), nullable=False)
    url          = db.Column(db.String(500), nullable=False)
    created_at   = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    added_by     = db.relationship('User', foreign_keys=[added_by_id])


@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


def ensure_schema_upgrades():
    """Add any newly-introduced columns to an existing database in place,
    so upgrading the code doesn't wipe or break existing data. Safe to call
    on every startup — it no-ops once columns already exist. Works against
    both SQLite (local dev) and Postgres (production), since the two use
    different type names for the same concept (e.g. DATETIME vs TIMESTAMP)."""
    from sqlalchemy import text, inspect
    inspector = inspect(db.engine)
    is_postgres = db.engine.dialect.name == 'postgresql'
    tables_needed = {
        'users': {
            'reset_token': 'VARCHAR(64)',
            'reset_token_expires': 'TIMESTAMP' if is_postgres else 'DATETIME',
            'organization_id': 'INTEGER',
        },
        'courses': {
            'lecturer_invite_code': 'VARCHAR(10)',
        },
        'session_participants': {
            'last_seen': 'TIMESTAMP' if is_postgres else 'DATETIME',
        },
    }
    existing_tables = set(inspector.get_table_names())
    for table_name, needed in tables_needed.items():
        if table_name not in existing_tables:
            continue  # fresh database — db.create_all() will create it with all columns
        existing_cols = {c['name'] for c in inspector.get_columns(table_name)}
        for col_name, col_type in needed.items():
            if col_name not in existing_cols:
                try:
                    db.session.execute(text(f'ALTER TABLE {table_name} ADD COLUMN {col_name} {col_type}'))
                    db.session.commit()
                    print(f'[MIGRATION] Added missing column {table_name}.{col_name}')
                except Exception as e:
                    db.session.rollback()
                    print(f'[MIGRATION ERROR] Could not add column {col_name} to {table_name}: {e}')


# ── Helpers ───────────────────────────────────────────────────────────────────

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


BLOCKED_SUBMISSION_EXTENSIONS = {'exe', 'sh', 'bat', 'cmd', 'msi', 'dll', 'com', 'scr', 'php', 'js', 'py', 'jar', 'apk'}

def allowed_submission_file(filename):
    """Assignment submissions can be 'any format' per course requirements —
    but still block obviously dangerous executable/script types rather
    than accepting literally anything."""
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[1].lower()
    return ext not in BLOCKED_SUBMISSION_EXTENSIONS


def generate_enrolment_code(length=6):
    chars = string.ascii_uppercase + string.digits
    while True:
        code = ''.join(secrets.choice(chars) for _ in range(length))
        if not Course.query.filter_by(enrolment_code=code).first():
            return code


def generate_lecturer_invite_code(length=8):
    chars = string.ascii_uppercase + string.digits
    while True:
        code = ''.join(secrets.choice(chars) for _ in range(length))
        if not Course.query.filter_by(lecturer_invite_code=code).first():
            return code


def generate_org_join_code(length=7):
    chars = string.ascii_uppercase + string.digits
    while True:
        code = ''.join(secrets.choice(chars) for _ in range(length))
        if not Organization.query.filter_by(join_code=code).first():
            return code


def is_course_lecturer(course, user):
    """True if user is the course's original owner OR a co-lecturer who
    was invited in. Use this everywhere a route currently only checks
    course.lecturer_id == current_user.id, so co-lecturers get the same
    access as the owner (except deleting the course, which stays
    owner-only — see delete_course)."""
    if course.lecturer_id == user.id:
        return True
    return CourseCoLecturer.query.filter_by(course_id=course.id, lecturer_id=user.id).first() is not None


def get_or_create_lecturer_invite_code(course):
    if not course.lecturer_invite_code:
        course.lecturer_invite_code = generate_lecturer_invite_code()
        db.session.commit()
    return course.lecturer_invite_code


def _send_via_brevo(subject, to_email, html, text):
    """Send via Brevo's HTTPS API (port 443 — never blocked). Returns True
    on success, False on failure (caller decides whether to fall back)."""
    try:
        resp = requests.post(
            'https://api.brevo.com/v3/smtp/email',
            headers={
                'api-key': app.config['BREVO_API_KEY'],
                'Content-Type': 'application/json',
                'Accept': 'application/json',
            },
            json={
                'sender': {'name': MAIL_SENDER_NAME, 'email': app.config['BREVO_SENDER_EMAIL']},
                'to': [{'email': to_email}],
                'subject': subject,
                'htmlContent': html,
                'textContent': text,
                'replyTo': {'email': app.config['BREVO_SENDER_EMAIL']},
            },
            timeout=10,
        )
        if resp.status_code in (200, 201):
            return True
        print(f"[BREVO ERROR] {resp.status_code}: {resp.text[:300]}")
        return False
    except Exception as e:
        print(f"[BREVO ERROR] {e}")
        return False


def _send_via_smtp(subject, to_email, html, text):
    """Send via Flask-Mail/SMTP. Only works if outbound SMTP isn't blocked
    (i.e. NOT Render's free tier — see BREVO_API_KEY above for the
    free-tier-compatible alternative)."""
    try:
        msg = Message(
            subject=subject,
            recipients=[to_email],
            html=html,
            body=text,
            reply_to=app.config['MAIL_USERNAME'] or None,
        )
        mail.send(msg)
        return True
    except Exception as e:
        print(f"[MAIL ERROR] {e}")
        return False


def _send_mail_async(subject, to_email, html, text):
    """Send an email in a background task so a slow/unreachable mail server
    can never hang the request that triggered it. Always includes a
    plain-text part alongside the HTML — mail with only an HTML body is a
    common spam-filter signal.

    Tries Brevo's HTTP API first if configured (works on Render's free
    tier, where outbound SMTP ports are blocked entirely), then falls back
    to SMTP if Brevo isn't configured or fails.

    Uses socketio.start_background_task() rather than a raw threading.Thread.
    Flask-SocketIO picks the right underlying primitive (real OS thread,
    eventlet greenlet, or gevent greenlet) to match whatever async_mode /
    gunicorn worker class is actually running the app — a raw
    threading.Thread would fight with an eventlet worker and corrupt
    SQLAlchemy's connection-pool locks (RuntimeError: cannot notify on
    un-acquired lock)."""
    def worker():
        with app.app_context():
            if app.config['BREVO_API_KEY']:
                if _send_via_brevo(subject, to_email, html, text):
                    return
                print("[MAIL] Brevo send failed, falling back to SMTP...")
            _send_via_smtp(subject, to_email, html, text)
    socketio.start_background_task(worker)


def send_otp_email(user):
    """Send OTP to user's email. Falls back silently if mail is unconfigured,
    and never blocks the request — the actual send happens on a background
    thread so a slow SMTP server can't freeze the page."""
    otp = user.generate_otp()
    db.session.commit()
    if not app.config['MAIL_USERNAME'] and not app.config['BREVO_API_KEY']:
        # No email method configured — print OTP to console for local dev
        print(f"\n[DEV] OTP for {user.email}: {otp}\n")
        return
    html = f"""
    <div style="font-family:sans-serif;max-width:480px;margin:auto;">
      <h2 style="color:#1d4ed8;">ClassConnect</h2>
      <p>Hello {user.full_name},</p>
      <p>Your one-time verification code is:</p>
      <div style="font-size:2rem;font-weight:700;letter-spacing:.4rem;
                  background:#eff6ff;padding:1rem;border-radius:8px;
                  text-align:center;color:#1e3a8a;">{otp}</div>
      <p style="color:#6b7280;font-size:.85rem;">
        This code expires in {OTP_EXPIRY_MINUTES} minutes.<br>
        If you did not register on ClassConnect, please ignore this email.
      </p>
    </div>"""
    text = (
        f"ClassConnect\n\nHello {user.full_name},\n\n"
        f"Your one-time verification code is: {otp}\n\n"
        f"This code expires in {OTP_EXPIRY_MINUTES} minutes.\n"
        f"If you did not register on ClassConnect, please ignore this email."
    )
    _send_mail_async('ClassConnect – Your verification code', user.email, html, text)


def send_reset_email(user):
    """Email a password-reset link. Same non-blocking pattern as OTP email."""
    token = user.generate_reset_token()
    db.session.commit()
    reset_url = url_for('reset_password', token=token, _external=True)
    if not app.config['MAIL_USERNAME'] and not app.config['BREVO_API_KEY']:
        print(f"\n[DEV] Password reset link for {user.email}: {reset_url}\n")
        return
    html = f"""
    <div style="font-family:sans-serif;max-width:480px;margin:auto;">
      <h2 style="color:#1d4ed8;">ClassConnect</h2>
      <p>Hello {user.full_name},</p>
      <p>We received a request to reset your password. Click the button below to choose a new one:</p>
      <p style="text-align:center;margin:1.5rem 0;">
        <a href="{reset_url}" style="display:inline-block;background:#2563eb;color:white;
           padding:.75rem 1.5rem;border-radius:8px;text-decoration:none;font-weight:600;">
           Reset Password</a>
      </p>
      <p style="color:#6b7280;font-size:.85rem;">
        This link expires in {RESET_TOKEN_EXPIRY_MINUTES} minutes.<br>
        If you did not request a password reset, please ignore this email — your password will stay unchanged.
      </p>
    </div>"""
    text = (
        f"ClassConnect\n\nHello {user.full_name},\n\n"
        f"We received a request to reset your password. Open this link to choose a new one:\n"
        f"{reset_url}\n\n"
        f"This link expires in {RESET_TOKEN_EXPIRY_MINUTES} minutes.\n"
        f"If you did not request a password reset, please ignore this email — your password will stay unchanged."
    )
    _send_mail_async('ClassConnect – Reset your password', user.email, html, text)


def role_required(*roles):
    def decorator(f):
        from functools import wraps
        @wraps(f)
        def wrapped(*args, **kwargs):
            if not current_user.is_authenticated or current_user.role not in roles:
                abort(403)
            return f(*args, **kwargs)
        return wrapped
    return decorator


@app.context_processor
def inject_org_labels():
    """Makes `labels` (and the raw org_type) available in every template
    automatically — e.g. {{ labels.admin_role }} renders 'Lecturer' for a
    university or 'Facilitator' for a general organization, without every
    single route needing to pass it explicitly."""
    if current_user.is_authenticated and current_user.organization:
        return {'labels': get_org_labels(current_user.organization),
                'current_org': current_user.organization}
    return {'labels': ORG_LABELS['university'], 'current_org': None}


# ── Auth routes ───────────────────────────────────────────────────────────────

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('index.html')


@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        full_name = request.form.get('full_name', '').strip()
        email     = request.form.get('email', '').strip().lower()
        password  = request.form.get('password', '')
        role      = request.form.get('role', 'student')
        id_number = request.form.get('id_number', '').strip().upper()
        org_mode  = request.form.get('org_mode', 'create')

        if role not in ('student', 'lecturer'):
            role = 'student'

        # Resolve which organization this account belongs to — either
        # joining an existing one via its code, or standing up a brand
        # new one (this person becomes its first member).
        organization = None
        if org_mode == 'join':
            org_code = request.form.get('org_join_code', '').strip().upper()
            organization = Organization.query.filter_by(join_code=org_code).first()
            if not organization:
                flash('That organization join code was not found. Check it and try again.', 'danger')
                return redirect(url_for('register'))
        else:
            org_name = request.form.get('org_name', '').strip()
            org_type = request.form.get('org_type', 'university')
            if org_type not in ('university', 'organization'):
                org_type = 'university'
            if not org_name:
                flash('Enter a name for your university or organization.', 'danger')
                return redirect(url_for('register'))
            organization = Organization(name=org_name, org_type=org_type,
                                        join_code=generate_org_join_code())
            db.session.add(organization)
            db.session.flush()   # get organization.id before we reference it below

        id_label = get_org_labels(organization)['id_field']

        if not full_name or not email or not password or not id_number:
            flash(f'All fields are required, including your {id_label}.', 'danger')
            return redirect(url_for('register'))
        if User.query.filter_by(email=email).first():
            flash('An account with that email already exists.', 'danger')
            return redirect(url_for('register'))
        if User.query.filter_by(id_number=id_number).first():
            flash(f'That {id_label} is already registered to another account.', 'danger')
            return redirect(url_for('register'))

        user = User(full_name=full_name, email=email, role=role, id_number=id_number,
                    organization_id=organization.id)
        user.set_password(password)
        db.session.add(user)
        db.session.flush()   # get user.id before commit
        send_otp_email(user)

        flash('Account created! Check your email for a 6-digit verification code.', 'success')
        return redirect(url_for('verify_otp', user_id=user.id))

    return render_template('register.html')


@app.route('/verify/<int:user_id>', methods=['GET', 'POST'])
def verify_otp(user_id):
    user = db.get_or_404(User, user_id)
    if user.email_verified:
        flash('Email already verified. Please log in.', 'info')
        return redirect(url_for('login'))

    if request.method == 'POST':
        code = request.form.get('otp', '').strip()
        if user.verify_otp(code):
            user.email_verified = True
            user.otp = None
            user.otp_expires = None
            db.session.commit()
            flash('Email verified! You can now log in.', 'success')
            return redirect(url_for('login'))
        else:
            flash('Invalid or expired code. Please try again or request a new one.', 'danger')

    return render_template('verify_otp.html', user=user)


@app.route('/verify/<int:user_id>/resend', methods=['POST'])
def resend_otp(user_id):
    user = db.get_or_404(User, user_id)
    if user.email_verified:
        return redirect(url_for('login'))
    send_otp_email(user)
    flash('A new code has been sent to your email.', 'info')
    return redirect(url_for('verify_otp', user_id=user.id))


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        identifier = request.form.get('identifier', '').strip()
        password   = request.form.get('password', '')

        # Allow logging in with either the registered email address or the
        # unique Matric Number (student) / Staff ID (lecturer).
        user = User.query.filter_by(email=identifier.lower()).first()
        if not user and identifier:
            user = User.query.filter_by(id_number=identifier.upper()).first()

        if user and user.check_password(password):
            if not user.email_verified:
                flash('Please verify your email first. Check your inbox for the OTP.', 'danger')
                return redirect(url_for('verify_otp', user_id=user.id))
            login_user(user)
            redirect_target = _complete_pending_join()
            return redirect(redirect_target or url_for('dashboard'))
        flash('Invalid email or password.', 'danger')
    return render_template('login.html')


def _complete_pending_join():
    """After logging in, if the person arrived via a /join/<code> or
    /courses/join-as-lecturer/<code> link, finish that action now and
    return where they should land — or None to fall back to the dashboard."""
    join_code = session.pop('pending_join_code', None)
    if join_code and current_user.role == 'student':
        course = Course.query.filter_by(enrolment_code=join_code).first()
        if course:
            already = _enrol_student(course, current_user)
            flash('You are already enrolled in this course.' if already
                  else f'Enrolled in {course.course_code} – {course.title}.', 'info' if already else 'success')
            return url_for('view_course', course_id=course.id)

    colecturer_code = session.pop('pending_colecturer_code', None)
    if colecturer_code and current_user.role == 'lecturer':
        course = Course.query.filter_by(lecturer_invite_code=colecturer_code).first()
        if course and course.lecturer_id != current_user.id:
            if not CourseCoLecturer.query.filter_by(course_id=course.id, lecturer_id=current_user.id).first():
                db.session.add(CourseCoLecturer(course_id=course.id, lecturer_id=current_user.id))
                db.session.commit()
                flash(f'You now have full lecturer access to "{course.title}".', 'success')
            return url_for('view_course', course_id=course.id)
    return None


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))


@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        user = User.query.filter_by(email=email).first()
        # Always show the same message whether or not the account exists,
        # so this form can't be used to check which emails are registered.
        if user:
            send_reset_email(user)
        flash('If an account exists for that email, a reset link has been sent.', 'info')
        return redirect(url_for('login'))
    return render_template('forgot_password.html')


@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    user = User.query.filter_by(reset_token=token).first()
    if not user or not user.verify_reset_token(token):
        flash('That reset link is invalid or has expired. Please request a new one.', 'danger')
        return redirect(url_for('forgot_password'))

    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm  = request.form.get('confirm_password', '')
        if len(password) < 6:
            flash('Password must be at least 6 characters.', 'danger')
            return render_template('reset_password.html', token=token)
        if password != confirm:
            flash('Passwords do not match.', 'danger')
            return render_template('reset_password.html', token=token)
        user.set_password(password)
        user.clear_reset_token()
        db.session.commit()
        flash('Password updated! You can now log in.', 'success')
        return redirect(url_for('login'))

    return render_template('reset_password.html', token=token)


# ── Dashboard ─────────────────────────────────────────────────────────────────

@app.route('/dashboard')
@login_required
def dashboard():
    if current_user.role == 'lecturer':
        courses = Course.query.filter_by(lecturer_id=current_user.id).all()
        return render_template('lecturer_dashboard.html', courses=courses)
    elif current_user.role == 'student':
        enrolments = Enrolment.query.filter_by(student_id=current_user.id).all()
        return render_template('student_dashboard.html', courses=[e.course for e in enrolments])
    else:
        users            = User.query.order_by(User.created_at.desc()).all()
        courses          = Course.query.all()
        total_sessions   = ClassSession.query.count()
        total_attendance = AttendanceRecord.query.count()
        return render_template('admin_dashboard.html', users=users, courses=courses,
                               total_sessions=total_sessions, total_attendance=total_attendance)


# ── Course management ─────────────────────────────────────────────────────────

@app.route('/courses/new', methods=['GET', 'POST'])
@login_required
@role_required('lecturer')
def create_course():
    if request.method == 'POST':
        course_code = request.form.get('course_code', '').strip()
        title       = request.form.get('title', '').strip()
        description = request.form.get('description', '').strip()
        if not course_code or not title:
            flash('Course code and title are required.', 'danger')
            return redirect(url_for('create_course'))
        course = Course(course_code=course_code, title=title, description=description,
                        enrolment_code=generate_enrolment_code(), lecturer_id=current_user.id)
        db.session.add(course)
        db.session.commit()
        flash(f'Course created. Enrolment code: {course.enrolment_code}', 'success')
        return redirect(url_for('view_course', course_id=course.id))
    return render_template('create_course.html')


@app.route('/courses/<int:course_id>')
@login_required
def view_course(course_id):
    course = db.get_or_404(Course, course_id)
    if current_user.role == 'lecturer' and not is_course_lecturer(course, current_user):
        abort(403)
    if current_user.role == 'student':
        if not Enrolment.query.filter_by(student_id=current_user.id, course_id=course.id).first():
            abort(403)

    sessions    = ClassSession.query.filter_by(course_id=course.id).order_by(ClassSession.started_at.desc()).all()
    materials   = CourseMaterial.query.filter_by(course_id=course.id).order_by(CourseMaterial.upload_date.desc()).all()
    links       = CourseLink.query.filter_by(course_id=course.id).order_by(CourseLink.created_at.desc()).all()
    assignments = Assignment.query.filter_by(course_id=course.id).order_by(Assignment.created_at.desc()).all()
    my_attendance = None
    if current_user.role == 'student':
        my_attendance = (AttendanceRecord.query.join(ClassSession)
                         .filter(ClassSession.course_id == course.id,
                                 AttendanceRecord.student_id == current_user.id).count())
    lecturer_invite_code = None
    if current_user.role == 'lecturer' and course.lecturer_id == current_user.id:
        lecturer_invite_code = get_or_create_lecturer_invite_code(course)
    return render_template('course_detail.html', course=course, sessions=sessions,
                           materials=materials, links=links, assignments=assignments,
                           my_attendance=my_attendance, co_lecturers=course.co_lecturers,
                           lecturer_invite_code=lecturer_invite_code,
                           is_owner=(course.lecturer_id == current_user.id))


@app.route('/courses/<int:course_id>/delete', methods=['POST'])
@login_required
@role_required('lecturer')
def delete_course(course_id):
    course = db.get_or_404(Course, course_id)
    if course.lecturer_id != current_user.id:
        abort(403)
    title = course.title
    # Collect uploaded material file paths before the DB rows (and their
    # cascade) are gone, so we can also clean up the actual files on disk.
    stored_filenames = [m.stored_filename for m in course.materials]
    for a in course.assignments:
        stored_filenames += [s.stored_filename for s in a.submissions]
    db.session.delete(course)  # cascades to enrolments, sessions, materials, links, co-lecturers, assignments (see model relationships)
    db.session.commit()
    for fname in stored_filenames:
        try:
            path = os.path.join(app.config['UPLOAD_FOLDER'], fname)
            if os.path.isfile(path):
                os.remove(path)
        except Exception as e:
            print(f'[DELETE_COURSE] Could not remove file {fname}: {e}')
    flash(f'Course "{title}" and all its data have been deleted.', 'success')
    return redirect(url_for('dashboard'))


def _enrol_student(course, student):
    """Shared enrolment logic used by both the manual code-entry form and
    the fast-join link. Returns (already_enrolled: bool)."""
    if Enrolment.query.filter_by(student_id=student.id, course_id=course.id).first():
        return True
    db.session.add(Enrolment(student_id=student.id, course_id=course.id))
    db.session.commit()
    return False


@app.route('/courses/enrol', methods=['POST'])
@login_required
@role_required('student')
def enrol_course():
    code   = request.form.get('enrolment_code', '').strip().upper()
    course = Course.query.filter_by(enrolment_code=code).first()
    if not course:
        flash('Invalid enrolment code.', 'danger')
        return redirect(url_for('dashboard'))
    already = _enrol_student(course, current_user)
    flash('You are already enrolled in this course.' if already
          else f'Enrolled in {course.course_code} – {course.title}.', 'info' if already else 'success')
    return redirect(url_for('view_course', course_id=course.id))


@app.route('/join/<code>')
def join_via_link(code):
    """Fast enrolment link — a student can just click a shared link instead
    of typing the course code manually. If not logged in, sends them to
    register/login first and comes back here automatically afterward."""
    course = Course.query.filter_by(enrolment_code=code.strip().upper()).first()
    if not course:
        flash('That course link is invalid or has expired.', 'danger')
        return redirect(url_for('index'))
    if not current_user.is_authenticated:
        session['pending_join_code'] = code.strip().upper()
        flash('Log in or create a student account to join this course.', 'info')
        return redirect(url_for('login'))
    if current_user.role != 'student':
        flash('Only student accounts can join a course this way.', 'danger')
        return redirect(url_for('dashboard'))
    already = _enrol_student(course, current_user)
    flash('You are already enrolled in this course.' if already
          else f'Enrolled in {course.course_code} – {course.title}.', 'info' if already else 'success')
    return redirect(url_for('view_course', course_id=course.id))


@app.route('/courses/join-as-lecturer/<code>')
def join_as_lecturer(code):
    """A different lecturer uses this link to become a co-lecturer on
    someone else's course, gaining full lecturer access to it (except
    deleting the course)."""
    course = Course.query.filter_by(lecturer_invite_code=code.strip().upper()).first()
    if not course:
        flash('That invite link is invalid or has expired.', 'danger')
        return redirect(url_for('index'))
    if not current_user.is_authenticated:
        session['pending_colecturer_code'] = code.strip().upper()
        flash('Log in with your lecturer account to accept this invite.', 'info')
        return redirect(url_for('login'))
    if current_user.role != 'lecturer':
        flash('Only lecturer accounts can accept this invite.', 'danger')
        return redirect(url_for('dashboard'))
    if course.lecturer_id == current_user.id:
        flash('This is already your own course.', 'info')
        return redirect(url_for('view_course', course_id=course.id))
    if not CourseCoLecturer.query.filter_by(course_id=course.id, lecturer_id=current_user.id).first():
        db.session.add(CourseCoLecturer(course_id=course.id, lecturer_id=current_user.id))
        db.session.commit()
        flash(f'You now have full lecturer access to "{course.title}".', 'success')
    return redirect(url_for('view_course', course_id=course.id))


@app.route('/courses/<int:course_id>/co-lecturers/<int:colecturer_id>/remove', methods=['POST'])
@login_required
@role_required('lecturer')
def remove_co_lecturer(course_id, colecturer_id):
    course = db.get_or_404(Course, course_id)
    if course.lecturer_id != current_user.id:  # only the original owner can revoke access
        abort(403)
    entry = db.get_or_404(CourseCoLecturer, colecturer_id)
    if entry.course_id != course.id:
        abort(404)
    db.session.delete(entry)
    db.session.commit()
    flash('Co-lecturer access removed.', 'success')
    return redirect(url_for('view_course', course_id=course.id))


# ── Sessions & attendance ─────────────────────────────────────────────────────

@app.route('/courses/<int:course_id>/sessions/start', methods=['POST'])
@login_required
@role_required('lecturer')
def start_session(course_id):
    course = db.get_or_404(Course, course_id)
    if not is_course_lecturer(course, current_user):
        abort(403)
    active = ClassSession.query.filter_by(course_id=course.id, status='active').first()
    if active:
        flash('A session is already active for this course.', 'info')
        return redirect(url_for('session_room', session_id=active.id))
    topic = request.form.get('topic', '').strip() or f'{course.course_code} session'
    s = ClassSession(course_id=course.id, topic=topic, status='active')
    db.session.add(s)
    db.session.commit()
    return redirect(url_for('session_room', session_id=s.id))


@app.route('/sessions/<int:session_id>')
@login_required
def session_room(session_id):
    session_obj = db.get_or_404(ClassSession, session_id)
    course      = session_obj.course

    if current_user.role == 'lecturer' and not is_course_lecturer(course, current_user):
        abort(403)
    if current_user.role == 'student':
        if not Enrolment.query.filter_by(student_id=current_user.id, course_id=course.id).first():
            abort(403)
        existing_participant = SessionParticipant.query.filter_by(
            session_id=session_obj.id, student_id=current_user.id).first()
        if existing_participant:
            if not existing_participant.still_connected:
                existing_participant.still_connected = True
                db.session.commit()
        elif session_obj.status == 'active':
            db.session.add(SessionParticipant(session_id=session_obj.id,
                                               student_id=current_user.id))
            db.session.commit()

    participants  = SessionParticipant.query.filter_by(session_id=session_obj.id).all()
    attendance    = AttendanceRecord.query.filter_by(session_id=session_obj.id).all()
    attendance_ids = {a.student_id for a in attendance}
    chat_history = (SessionChatMessage.query.filter_by(session_id=session_obj.id)
                    .order_by(SessionChatMessage.created_at).all())
    chat_history_json = [{
        'id': m.id, 'user_id': m.user_id, 'user_name': m.user.full_name, 'role': m.user.role,
        'text': m.text, 'ts': m.created_at.strftime('%H:%M'), 'edited': m.edited_at is not None,
    } for m in chat_history]

    return render_template('session_room.html', session=session_obj, course=course,
                           participants=participants, attendance_ids=attendance_ids,
                           chat_history_json=chat_history_json)


@app.route('/api/turn-credentials')
@login_required
def turn_credentials():
    """Return fresh ICE servers (STUN + TURN) for the browser's
    RTCPeerConnection. TURN is essential when the two peers are on
    different networks with restrictive NATs (e.g. one on mobile data,
    one on wifi) — STUN alone frequently can't establish that connection,
    which shows up as signalling working (peers see each other in the
    room) but video never actually connecting.

    Uses the free Open Relay TURN service (metered.ca/tools/openrelay) —
    sign up free, create an app, get your API key. Set TURN_API_KEY and
    TURN_APP_NAME (the app name you chose at signup — it becomes part of
    the API URL: https://<appname>.metered.live/...) as env vars. The key
    is only ever used server-side here, never sent to the browser. Falls
    back to Google's public STUN-only servers if not configured, which
    still works fine for peers on the same network."""
    fallback = {'iceServers': [
        {'urls': 'stun:stun.l.google.com:19302'},
        {'urls': 'stun:stun1.l.google.com:19302'},
    ]}
    api_key  = os.environ.get('TURN_API_KEY', '')
    app_name = os.environ.get('TURN_APP_NAME', '')
    if not api_key or not app_name:
        return jsonify(fallback)
    try:
        resp = requests.get(
            f'https://{app_name}.metered.live/api/v1/turn/credentials',
            params={'apiKey': api_key},
            timeout=5,
        )
        if resp.status_code == 200:
            return jsonify({'iceServers': resp.json()})
        print(f'[TURN ERROR] {resp.status_code}: {resp.text[:200]}')
    except Exception as e:
        print(f'[TURN ERROR] {e}')
    return jsonify(fallback)


def _currently_connected_participants(session_id):
    """The real, self-healing definition of 'connected right now': not just
    the still_connected flag (which can get stuck if a disconnect event was
    ever missed), but also a recent heartbeat. Someone whose last heartbeat
    is older than PRESENCE_STALE_SECONDS is treated as gone, regardless of
    what the flag says."""
    cutoff = datetime.now(timezone.utc) - timedelta(seconds=PRESENCE_STALE_SECONDS)
    return (SessionParticipant.query
            .filter_by(session_id=session_id, still_connected=True)
            .filter(SessionParticipant.last_seen >= cutoff)
            .all())


@app.route('/sessions/<int:session_id>/participants')
@login_required
def session_participants_json(session_id):
    session_obj  = db.get_or_404(ClassSession, session_id)
    participants = _currently_connected_participants(session_obj.id)
    attendance_ids = {a.student_id for a in
                      AttendanceRecord.query.filter_by(session_id=session_obj.id).all()}
    data = [{'student_id': p.student_id, 'name': p.student.full_name,
              'id_number': p.student.id_number or '—',
              'joined_at': p.joined_at.strftime('%H:%M:%S'),
              'present': p.student_id in attendance_ids} for p in participants]
    return jsonify({'status': session_obj.status, 'participants': data})


@app.route('/sessions/<int:session_id>/take-attendance', methods=['POST'])
@login_required
@role_required('lecturer')
def take_attendance(session_id):
    session_obj = db.get_or_404(ClassSession, session_id)
    if not is_course_lecturer(session_obj.course, current_user):
        abort(403)
    participants = _currently_connected_participants(session_obj.id)
    already = {a.student_id for a in
               AttendanceRecord.query.filter_by(session_id=session_obj.id).all()}
    count = 0
    for p in participants:
        if p.student_id not in already:
            db.session.add(AttendanceRecord(session_id=session_obj.id,
                                             student_id=p.student_id, status='present'))
            count += 1
    db.session.commit()
    flash(f'Attendance captured: {count} new student(s) marked present.', 'success')
    return redirect(url_for('session_room', session_id=session_obj.id))


@app.route('/sessions/<int:session_id>/close', methods=['POST'])
@login_required
@role_required('lecturer')
def close_session(session_id):
    session_obj = db.get_or_404(ClassSession, session_id)
    if not is_course_lecturer(session_obj.course, current_user):
        abort(403)
    session_obj.status    = 'closed'
    session_obj.closed_at = datetime.now(timezone.utc)
    db.session.commit()
    flash('Session closed.', 'info')
    return redirect(url_for('view_course', course_id=session_obj.course_id))


@app.route('/sessions/<int:session_id>/report')
@login_required
def session_report(session_id):
    session_obj = db.get_or_404(ClassSession, session_id)
    course      = session_obj.course
    if current_user.role == 'lecturer' and not is_course_lecturer(course, current_user):
        abort(403)
    if current_user.role == 'student':
        if not Enrolment.query.filter_by(student_id=current_user.id, course_id=course.id).first():
            abort(403)
    records = AttendanceRecord.query.filter_by(session_id=session_obj.id).all()
    return render_template('session_report.html', session=session_obj,
                           course=course, records=records)


def _build_attendance_matrix(course):
    """Returns (sessions, students, matrix) where matrix[student_id][session_id] = bool present."""
    sessions = ClassSession.query.filter_by(course_id=course.id).order_by(ClassSession.started_at).all()
    students = ([e.student for e in
                 Enrolment.query.filter_by(course_id=course.id).join(User).order_by(User.full_name).all()])
    records = (AttendanceRecord.query.join(ClassSession)
              .filter(ClassSession.course_id == course.id).all())
    present_set = {(r.student_id, r.session_id) for r in records}
    matrix = {s.id: {sess.id: (s.id, sess.id) in present_set for sess in sessions} for s in students}
    return sessions, students, matrix


@app.route('/courses/<int:course_id>/attendance/export/<fmt>')
@login_required
@role_required('lecturer')
def export_attendance(course_id, fmt):
    course = db.get_or_404(Course, course_id)
    if not is_course_lecturer(course, current_user):
        abort(403)
    if fmt not in ('pdf', 'docx'):
        abort(404)

    sessions, students, matrix = _build_attendance_matrix(course)
    safe_code = secure_filename(course.course_code) or 'course'

    if fmt == 'pdf':
        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                                topMargin=1.2*cm, bottomMargin=1.2*cm, leftMargin=1.2*cm, rightMargin=1.2*cm)
        styles = getSampleStyleSheet()
        elements = [
            Paragraph(f"Attendance Register — {course.course_code}: {course.title}", styles['Title']),
            Paragraph(f"Generated {datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M UTC')}", styles['Normal']),
            Spacer(1, 12),
        ]
        header = ['Student', 'Matric No.'] + [s.started_at.strftime('%d/%m') for s in sessions] + ['Total']
        rows = [header]
        for stu in students:
            present_count = sum(1 for sess in sessions if matrix[stu.id][sess.id])
            row = [stu.full_name, stu.id_number or '—']
            row += ['✓' if matrix[stu.id][sess.id] else '—' for sess in sessions]
            row.append(f"{present_count}/{len(sessions)}")
            rows.append(row)
        if not students:
            rows.append(['No enrolled students yet.'] + [''] * (len(sessions) + 1))
        table = Table(rows, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a8a')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
            ('ALIGN', (2, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        elements.append(table)
        doc.build(elements)
        buf.seek(0)
        return (buf.read(), 200, {
            'Content-Type': 'application/pdf',
            'Content-Disposition': f'attachment; filename="attendance_{safe_code}.pdf"',
        })

    else:  # docx
        buf = BytesIO()
        doc = DocxDocument()
        doc.add_heading(f"Attendance Register — {course.course_code}: {course.title}", level=1)
        p = doc.add_paragraph(f"Generated {datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M UTC')}")
        p.runs[0].font.size = Pt(9)

        cols = 2 + len(sessions) + 1
        table = doc.add_table(rows=1, cols=cols)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Student'
        hdr[1].text = 'Matric No.'
        for i, sess in enumerate(sessions):
            hdr[2 + i].text = sess.started_at.strftime('%d/%m')
        hdr[-1].text = 'Total'

        for stu in students:
            present_count = sum(1 for sess in sessions if matrix[stu.id][sess.id])
            row = table.add_row().cells
            row[0].text = stu.full_name
            row[1].text = stu.id_number or '—'
            for i, sess in enumerate(sessions):
                row[2 + i].text = '✓' if matrix[stu.id][sess.id] else '—'
            row[-1].text = f"{present_count}/{len(sessions)}"
        if not students:
            doc.add_paragraph('No enrolled students yet.')

        doc.save(buf)
        buf.seek(0)
        return (buf.read(), 200, {
            'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'Content-Disposition': f'attachment; filename="attendance_{safe_code}.docx"',
        })


# ── Materials ─────────────────────────────────────────────────────────────────

@app.route('/courses/<int:course_id>/materials/upload', methods=['POST'])
@login_required
@role_required('lecturer')
def upload_material(course_id):
    course = db.get_or_404(Course, course_id)
    if not is_course_lecturer(course, current_user):
        abort(403)
    title = request.form.get('title', '').strip()
    file  = request.files.get('file')
    if not title or not file or file.filename == '':
        flash('A title and a file are required.', 'danger')
        return redirect(url_for('view_course', course_id=course.id))
    if not allowed_file(file.filename):
        flash('Only PDF, Word, and PowerPoint files are allowed.', 'danger')
        return redirect(url_for('view_course', course_id=course.id))
    original = secure_filename(file.filename)
    stored   = f"{secrets.token_hex(8)}_{original}"
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], stored))
    db.session.add(CourseMaterial(course_id=course.id, uploader_id=current_user.id,
                                  title=title, filename=original, stored_filename=stored))
    db.session.commit()
    flash('Material uploaded.', 'success')
    return redirect(url_for('view_course', course_id=course.id))


@app.route('/materials/<int:material_id>/download')
@login_required
def download_material(material_id):
    material = db.get_or_404(CourseMaterial, material_id)
    course   = material.course
    if current_user.role == 'lecturer' and not is_course_lecturer(course, current_user):
        abort(403)
    if current_user.role == 'student':
        if not Enrolment.query.filter_by(student_id=current_user.id, course_id=course.id).first():
            abort(403)
    return send_from_directory(app.config['UPLOAD_FOLDER'], material.stored_filename,
                               as_attachment=True, download_name=material.filename)


@app.route('/materials/<int:material_id>/delete', methods=['POST'])
@login_required
@role_required('lecturer')
def delete_material(material_id):
    material = db.get_or_404(CourseMaterial, material_id)
    if not is_course_lecturer(material.course, current_user):
        abort(403)
    course_id = material.course_id
    stored_filename = material.stored_filename
    db.session.delete(material)
    db.session.commit()
    try:
        path = os.path.join(app.config['UPLOAD_FOLDER'], stored_filename)
        if os.path.isfile(path):
            os.remove(path)
    except Exception as e:
        print(f'[DELETE_MATERIAL] Could not remove file {stored_filename}: {e}')
    flash('Material deleted.', 'success')
    return redirect(url_for('view_course', course_id=course_id))


@app.route('/courses/<int:course_id>/links/add', methods=['POST'])
@login_required
@role_required('lecturer')
def add_course_link(course_id):
    course = db.get_or_404(Course, course_id)
    if not is_course_lecturer(course, current_user):
        abort(403)
    title = request.form.get('title', '').strip()
    url_  = request.form.get('url', '').strip()
    if not title or not url_:
        flash('A title and a URL are required.', 'danger')
        return redirect(url_for('view_course', course_id=course.id))
    if not (url_.startswith('http://') or url_.startswith('https://')):
        url_ = 'https://' + url_
    db.session.add(CourseLink(course_id=course.id, added_by_id=current_user.id, title=title, url=url_))
    db.session.commit()
    flash('Link added.', 'success')
    return redirect(url_for('view_course', course_id=course.id))


@app.route('/links/<int:link_id>/delete', methods=['POST'])
@login_required
@role_required('lecturer')
def delete_course_link(link_id):
    link = db.get_or_404(CourseLink, link_id)
    if not is_course_lecturer(link.course, current_user):
        abort(403)
    course_id = link.course_id
    db.session.delete(link)
    db.session.commit()
    flash('Link removed.', 'success')
    return redirect(url_for('view_course', course_id=course_id))


# ── Assignments ──────────────────────────────────────────────────────────────

@app.route('/courses/<int:course_id>/assignments/new', methods=['POST'])
@login_required
@role_required('lecturer')
def create_assignment(course_id):
    course = db.get_or_404(Course, course_id)
    if not is_course_lecturer(course, current_user):
        abort(403)
    title        = request.form.get('title', '').strip()
    instructions = request.form.get('instructions', '').strip()
    due_date_str = request.form.get('due_date', '').strip()
    if not title:
        flash('An assignment title is required.', 'danger')
        return redirect(url_for('view_course', course_id=course.id))
    due_date = None
    if due_date_str:
        try:
            due_date = datetime.strptime(due_date_str, '%Y-%m-%dT%H:%M')
        except ValueError:
            pass
    db.session.add(Assignment(course_id=course.id, created_by=current_user.id,
                              title=title, instructions=instructions, due_date=due_date))
    db.session.commit()
    flash('Assignment posted.', 'success')
    return redirect(url_for('view_course', course_id=course.id))


@app.route('/assignments/<int:assignment_id>')
@login_required
def view_assignment(assignment_id):
    assignment = db.get_or_404(Assignment, assignment_id)
    course = assignment.course
    if current_user.role == 'lecturer' and not is_course_lecturer(course, current_user):
        abort(403)
    if current_user.role == 'student':
        if not Enrolment.query.filter_by(student_id=current_user.id, course_id=course.id).first():
            abort(403)

    my_submission = None
    submissions = []
    if current_user.role == 'student':
        my_submission = AssignmentSubmission.query.filter_by(
            assignment_id=assignment.id, student_id=current_user.id).first()
    else:
        submissions = (AssignmentSubmission.query.filter_by(assignment_id=assignment.id)
                       .order_by(AssignmentSubmission.submitted_at.desc()).all())
    return render_template('assignment_detail.html', assignment=assignment, course=course,
                           my_submission=my_submission, submissions=submissions)


@app.route('/assignments/<int:assignment_id>/submit', methods=['POST'])
@login_required
@role_required('student')
def submit_assignment(assignment_id):
    assignment = db.get_or_404(Assignment, assignment_id)
    course = assignment.course
    if not Enrolment.query.filter_by(student_id=current_user.id, course_id=course.id).first():
        abort(403)
    file = request.files.get('file')
    if not file or file.filename == '':
        flash('Choose a file to submit.', 'danger')
        return redirect(url_for('view_assignment', assignment_id=assignment.id))
    if not allowed_submission_file(file.filename):
        flash('That file type is not allowed for submissions.', 'danger')
        return redirect(url_for('view_assignment', assignment_id=assignment.id))

    original = secure_filename(file.filename)
    stored   = f"{secrets.token_hex(8)}_{original}"
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    file.save(os.path.join(app.config['UPLOAD_FOLDER'], stored))

    # One submission per student — resubmitting replaces the previous file,
    # so there's always exactly one clearly-identified submission per
    # student for the lecturer to grade.
    existing = AssignmentSubmission.query.filter_by(
        assignment_id=assignment.id, student_id=current_user.id).first()
    if existing:
        old_path = os.path.join(app.config['UPLOAD_FOLDER'], existing.stored_filename)
        if os.path.isfile(old_path):
            try: os.remove(old_path)
            except Exception: pass
        existing.filename = original
        existing.stored_filename = stored
        existing.submitted_at = datetime.now(timezone.utc)
    else:
        db.session.add(AssignmentSubmission(assignment_id=assignment.id, student_id=current_user.id,
                                            filename=original, stored_filename=stored))
    db.session.commit()
    flash('Assignment submitted.', 'success')
    return redirect(url_for('view_assignment', assignment_id=assignment.id))


@app.route('/assignments/<int:assignment_id>/delete', methods=['POST'])
@login_required
@role_required('lecturer')
def delete_assignment(assignment_id):
    assignment = db.get_or_404(Assignment, assignment_id)
    if not is_course_lecturer(assignment.course, current_user):
        abort(403)
    course_id = assignment.course_id
    stored_filenames = [s.stored_filename for s in assignment.submissions]
    db.session.delete(assignment)  # cascades to submissions
    db.session.commit()
    for fname in stored_filenames:
        try:
            path = os.path.join(app.config['UPLOAD_FOLDER'], fname)
            if os.path.isfile(path):
                os.remove(path)
        except Exception as e:
            print(f'[DELETE_ASSIGNMENT] Could not remove file {fname}: {e}')
    flash('Assignment deleted.', 'success')
    return redirect(url_for('view_course', course_id=course_id))


@app.route('/submissions/<int:submission_id>/download')
@login_required
def download_submission(submission_id):
    sub = db.get_or_404(AssignmentSubmission, submission_id)
    course = sub.assignment.course
    is_lecturer_of_course = current_user.role == 'lecturer' and is_course_lecturer(course, current_user)
    is_the_student = current_user.id == sub.student_id
    if not (is_lecturer_of_course or is_the_student):
        abort(403)
    return send_from_directory(app.config['UPLOAD_FOLDER'], sub.stored_filename,
                               as_attachment=True, download_name=sub.filename)


# ── Admin ─────────────────────────────────────────────────────────────────────

@app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
@login_required
@role_required('admin')
def delete_user(user_id):
    if user_id == current_user.id:
        flash("You can't delete your own account.", 'danger')
        return redirect(url_for('dashboard'))
    user = db.get_or_404(User, user_id)
    db.session.delete(user)
    db.session.commit()
    flash('User removed.', 'info')
    return redirect(url_for('dashboard'))


# ── WebRTC signalling (SocketIO events) ───────────────────────────────────────
# Each class session gets its own SocketIO room named "session_<id>".
# The server only relays signalling messages (offer/answer/ICE candidates)
# between peers; it never handles media — all video/audio flows peer-to-peer
# via WebRTC using Google's free STUN servers.

# Maps a live socket connection (sid) to who it belongs to. Needed because
# Socket.IO's 'disconnect' event doesn't carry the original join payload —
# without this we'd have no reliable way to know who dropped when a phone
# gets locked, a tab is killed, or the network drops, none of which fire
# the client's 'beforeunload' cleanup in time (or at all, on mobile).
#
# IMPORTANT: this dict lives in process memory, and Render's free tier
# spins the whole process down after ~15 minutes idle (and restarts it on
# the next request). A restart wipes this dict, so relying on the
# disconnect-event path ALONE would leave anyone who was connected before
# the restart stuck marked "still connected" forever — nothing would ever
# fire a disconnect for them again. To make this self-healing regardless
# of restarts, PRESENCE_STALE_SECONDS below is the real source of truth:
# every connected client sends a periodic heartbeat, and anyone whose
# heartbeat has gone stale (whatever the reason — clean leave, crash,
# network drop, or a server restart in between) simply stops showing up,
# without needing to catch the exact moment they left.
sid_registry = {}
PRESENCE_STALE_SECONDS = 25   # heartbeat is sent every 10s client-side


def _touch_participant(session_id, user_id):
    """Mark a student as freshly seen — called on join and on every heartbeat."""
    sp = SessionParticipant.query.filter_by(session_id=session_id, student_id=user_id).first()
    if sp:
        sp.still_connected = True
        sp.last_seen = datetime.now(timezone.utc)
        db.session.commit()


def _mark_participant_disconnected(session_id, user_id):
    sp = SessionParticipant.query.filter_by(session_id=session_id, student_id=user_id).first()
    if sp:
        sp.still_connected = False
        db.session.commit()


@socketio.on('join-video-room')
def on_join_video(data):
    room      = f"session_{data['session_id']}"
    user_id   = data['user_id']
    user_name = data['user_name']
    join_room(room)
    sid_registry[request.sid] = {'session_id': data['session_id'], 'user_id': user_id, 'user_name': user_name}
    # Tell everyone else in the room that a new peer has arrived
    emit('peer-joined', {'user_id': user_id, 'user_name': user_name}, to=room, skip_sid=request.sid)  # noqa

    if current_user.is_authenticated and current_user.role == 'student':
        _touch_participant(data['session_id'], current_user.id)


@socketio.on('presence-heartbeat')
def on_presence_heartbeat(data):
    """Sent every ~10s by the client while the session page is open and
    connected. This is what actually keeps someone showing as 'connected'
    — not a one-time flag — so the roster self-heals if a disconnect was
    ever missed (including across a Render free-tier restart)."""
    if current_user.is_authenticated and current_user.role == 'student':
        _touch_participant(data['session_id'], current_user.id)


@socketio.on('leave-video-room')
def on_leave_video(data):
    room = f"session_{data['session_id']}"
    leave_room(room)
    emit('peer-left', {'user_id': data['user_id']}, to=room)
    _mark_participant_disconnected(data['session_id'], data['user_id'])
    sid_registry.pop(request.sid, None)


@socketio.on('disconnect')
def on_socket_disconnect():
    """Catches clean-ish disconnects for IMMEDIATE roster updates (nicer
    UX than waiting out the heartbeat timeout). The heartbeat timeout
    above is what guarantees correctness even when this never fires."""
    info = sid_registry.pop(request.sid, None)
    if not info:
        return
    room = f"session_{info['session_id']}"
    emit('peer-left', {'user_id': info['user_id']}, to=room)
    _mark_participant_disconnected(info['session_id'], info['user_id'])


@socketio.on('camera-state-changed')
def on_camera_state_changed(data):
    """Relayed so every remote peer can show an accurate placeholder
    (avatar + name) instead of a frozen/black video square when someone's
    camera is off — while their audio, if unmuted, keeps working
    independently. Relying on raw WebRTC track state for this is
    inconsistent across browsers, so the app signals it explicitly."""
    room = f"session_{data['session_id']}"
    emit('camera-state-changed', {
        'user_id': data['user_id'], 'camera_on': bool(data.get('camera_on'))
    }, to=room, skip_sid=request.sid)


@socketio.on('offer')
def on_offer(data):
    room = f"session_{data['session_id']}"
    emit('offer', data, to=room, skip_sid=request.sid)  # noqa


@socketio.on('answer')
def on_answer(data):
    room = f"session_{data['session_id']}"
    emit('answer', data, to=room, skip_sid=request.sid)  # noqa


@socketio.on('ice-candidate')
def on_ice_candidate(data):
    room = f"session_{data['session_id']}"
    emit('ice-candidate', data, to=room, skip_sid=request.sid)  # noqa


# ── Live-session moderation & interaction ─────────────────────────────────────

@socketio.on('raise-hand')
def on_raise_hand(data):
    """A student signals they want to ask a question. Relayed to everyone
    in the room (lecturer's UI shows it as an actionable list; other
    students just see who has their hand up, same as any classroom)."""
    room = f"session_{data['session_id']}"
    emit('hand-raised', {'user_id': data['user_id'], 'user_name': data['user_name']},
         to=room, skip_sid=request.sid)


@socketio.on('lower-hand')
def on_lower_hand(data):
    room = f"session_{data['session_id']}"
    emit('hand-lowered', {'user_id': data['user_id']}, to=room, skip_sid=request.sid)


@socketio.on('spotlight-student')
def on_spotlight_student(data):
    """Lecturer picks a student (usually one with a raised hand) to be
    visually featured/enlarged for everyone in the session. Sending
    user_id=null clears the spotlight. Lecturer-only."""
    if not current_user.is_authenticated or current_user.role != 'lecturer':
        return
    room = f"session_{data['session_id']}"
    emit('spotlight-changed', {'user_id': data.get('user_id')}, to=room)  # includes sender, for UI consistency


@socketio.on('force-mute')
def on_force_mute(data):
    """Lecturer mutes a specific student's mic. The server only relays this
    request — the student's own browser is what actually disables their
    track, same pattern used by Zoom/Meet/Teams host controls. Lecturer-only."""
    if not current_user.is_authenticated or current_user.role != 'lecturer':
        return
    room = f"session_{data['session_id']}"
    emit('force-mute', {'user_id': data['target_user_id']}, to=room, skip_sid=request.sid)


@socketio.on('force-unmute')
def on_force_unmute(data):
    """Lecturer asks a specific student to unmute. Lecturer-only."""
    if not current_user.is_authenticated or current_user.role != 'lecturer':
        return
    room = f"session_{data['session_id']}"
    emit('force-unmute', {'user_id': data['target_user_id']}, to=room, skip_sid=request.sid)


# ── Live chat ────────────────────────────────────────────────────────────────

@socketio.on('chat-message')
def on_chat_message(data):
    """Live-session chat — now persisted to the database (see
    SessionChatMessage), so re-entering an active session or just
    refreshing shows the full history instead of losing it. Broadcast to
    everyone in the room, including the sender, so every client renders
    from a single authoritative source."""
    if not current_user.is_authenticated:
        return
    session_id = data['session_id']
    room = f"session_{session_id}"
    text = (data.get('text') or '').strip()[:1000]  # hard cap, avoid abuse
    if not text:
        return

    msg = SessionChatMessage(session_id=session_id, user_id=current_user.id, text=text)
    db.session.add(msg)
    db.session.commit()

    emit('chat-message', {
        'id': msg.id,
        'user_id': current_user.id,
        'user_name': current_user.full_name,
        'role': current_user.role,
        'text': text,
        'ts': msg.created_at.strftime('%H:%M'),
    }, to=room)


@socketio.on('edit-chat-message')
def on_edit_chat_message(data):
    """Lets someone fix a typo in a message they sent earlier in the same
    live session. Only the original author can edit their own message —
    enforced here server-side, not just hidden in the UI."""
    if not current_user.is_authenticated:
        return
    msg = db.session.get(SessionChatMessage, data.get('message_id'))
    if not msg or msg.user_id != current_user.id:
        return   # not their message — silently ignore, no info leak either way
    new_text = (data.get('text') or '').strip()[:1000]
    if not new_text:
        return
    msg.text = new_text
    msg.edited_at = datetime.now(timezone.utc)
    db.session.commit()

    room = f"session_{msg.session_id}"
    emit('chat-message-edited', {
        'id': msg.id, 'text': msg.text, 'edited': True
    }, to=room)


# ── Student screen-share approval ───────────────────────────────────────────

@socketio.on('request-screen-share')
def on_request_screen_share(data):
    """A student asks to share their screen. Relayed to the lecturer only
    (broadcast to the room; every non-lecturer client's listener just
    ignores it — simplest way to reach 'the lecturer' without tracking
    individual socket IDs per role)."""
    if not current_user.is_authenticated or current_user.role != 'student':
        return
    room = f"session_{data['session_id']}"
    emit('screen-share-requested', {
        'user_id': current_user.id, 'user_name': current_user.full_name
    }, to=room, skip_sid=request.sid)


@socketio.on('respond-screen-share')
def on_respond_screen_share(data):
    """Lecturer approves or denies a student's screen-share request.
    Lecturer-only. On approval, also sets that student as the spotlight —
    the whole point of asking was to be seen presenting something."""
    if not current_user.is_authenticated or current_user.role != 'lecturer':
        return
    room = f"session_{data['session_id']}"
    approved = bool(data.get('approved'))
    target_id = data['target_user_id']
    emit('screen-share-response', {'approved': approved, 'user_id': target_id}, to=room)
    if approved:
        emit('spotlight-changed', {'user_id': target_id}, to=room)


@socketio.on('student-screen-share-ended')
def on_student_screen_share_ended(data):
    """A student's screen share ended (they stopped it, or the browser's
    native 'stop sharing' control fired). Clears the spotlight for
    everyone automatically — 'remove them from spotlight when they're
    done', without requiring the lecturer to do it manually."""
    if not current_user.is_authenticated:
        return
    room = f"session_{data['session_id']}"
    emit('spotlight-changed', {'user_id': None}, to=room)


@socketio.on('revoke-screen-share')
def on_revoke_screen_share(data):
    """Lecturer manually ends a student's screen share early. Lecturer-only."""
    if not current_user.is_authenticated or current_user.role != 'lecturer':
        return
    room = f"session_{data['session_id']}"
    emit('screen-share-revoked', {'user_id': data['target_user_id']}, to=room)
    emit('spotlight-changed', {'user_id': None}, to=room)


# ── Error handlers ────────────────────────────────────────────────────────────

@app.errorhandler(403)
def forbidden(e):
    return render_template('error.html', code=403,
                           message="You don't have permission to view this page."), 403

@app.errorhandler(404)
def not_found(e):
    return render_template('error.html', code=404, message='Page not found.'), 404


# ── CLI ───────────────────────────────────────────────────────────────────────

@app.cli.command('init-db')
def init_db():
    db.create_all()
    ensure_schema_upgrades()
    if not User.query.filter_by(role='admin').first():
        admin = User(full_name='System Administrator', email='admin@fuo.edu.ng',
                     role='admin', email_verified=True)
        admin.set_password('Admin@123')
        db.session.add(admin)
        db.session.commit()
        print('Created default admin -> email: admin@fuo.edu.ng  password: Admin@123')
    print('Database initialised.')


def bootstrap_db():
    """Create the instance folder + database + default admin, and apply any
    pending column migrations. Runs unconditionally at import time (not just
    under `if __name__ == '__main__'`) so this also works correctly when the
    app is started via gunicorn/eventlet in production, where this module is
    imported rather than executed directly."""
    with app.app_context():
        os.makedirs(os.path.join(basedir, 'instance'), exist_ok=True)
        db.create_all()
        ensure_schema_upgrades()

        # This deployment predates the multi-organization model — every
        # existing user needs to belong to SOME organization so their
        # courses/enrolments/sessions keep working coherently instead of
        # looking orphaned. Create one default university organization and
        # attach any user who doesn't already have one.
        orphaned = User.query.filter_by(organization_id=None).all()
        if orphaned:
            default_org = Organization.query.filter_by(name='Federal University Otuoke').first()
            if not default_org:
                default_org = Organization(name='Federal University Otuoke', org_type='university',
                                           join_code=generate_org_join_code())
                db.session.add(default_org)
                db.session.flush()
            for u in orphaned:
                u.organization_id = default_org.id
            db.session.commit()
            print(f'[MIGRATION] Assigned {len(orphaned)} pre-existing user(s) to default organization '
                  f'"{default_org.name}" (join code: {default_org.join_code})')

        if not User.query.filter_by(role='admin').first():
            admin = User(full_name='System Administrator', email='admin@fuo.edu.ng',
                         role='admin', email_verified=True)
            admin.set_password('Admin@123')
            db.session.add(admin)
            db.session.commit()
            print('Created default admin -> admin@fuo.edu.ng / Admin@123')
        # Critical: dispose of the connection pool used during this bootstrap.
        # If the app is imported once and then forked into multiple worker
        # processes (gunicorn "preload_app"), each fork would otherwise
        # inherit this pool's internal locks in whatever state they happened
        # to be in at fork time — which can permanently break with
        # "RuntimeError: cannot notify on un-acquired lock" the moment a
        # forked worker tries to use it. Disposing here forces every worker
        # to lazily open its own fresh connection instead of reusing this one.
        db.engine.dispose()


bootstrap_db()

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    socketio.run(app, debug=True, host='0.0.0.0', port=port, allow_unsafe_werkzeug=True)
